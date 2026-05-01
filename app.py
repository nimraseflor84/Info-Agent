"""
AUTODOC Info Agent — Hauptanwendung v3
Erweitert: Artikel mit Bildern, Tabellen, Publikationsformat, HTML-Export
"""

import json
import sqlite3
import threading
import webbrowser
import time
import os
import base64
import re
from datetime import datetime, date
from flask import Flask, render_template, jsonify, request
from apscheduler.schedulers.background import BackgroundScheduler
import feedparser
import anthropic
import requests
from urllib.parse import quote_plus
from concurrent.futures import ThreadPoolExecutor, as_completed

# ── Feed-Status (für Frontend-Polling) ────────────────────────────────────────
feed_status = {"running": False, "last_fetch": "—", "new_count": 0}

# ── Pfade ────────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")
SOURCES_FILE = os.path.join(BASE_DIR, "sources.json")
DB_FILE = os.path.join(BASE_DIR, "memory.db")

# ── Config ────────────────────────────────────────────────────────────────────
def load_config():
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_config(data):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# ── Quellliste laden (JSON mit Kommentaren tolerieren + DB-Quellen) ──────────
def load_sources():
    with open(SOURCES_FILE, "r", encoding="utf-8") as f:
        content = f.read()
    # Kommentarzeilen entfernen (nur Zeilen, die mit // beginnen — nicht // in URLs)
    content = re.sub(r'^\s*//[^\n]*\n?', '', content, flags=re.MULTILINE)
    sources = json.loads(content)
    # Eigene Quellen aus DB anhängen
    try:
        conn = get_db()
        custom = conn.execute("SELECT * FROM custom_sources WHERE active = 1").fetchall()
        conn.close()
        for cs in custom:
            cs = dict(cs)  # sqlite3.Row → dict (hat .get())
            # Subscriber-RSS als primäre URL verwenden wenn vorhanden
            feed_url = cs.get("subscriber_rss") or cs["url"]
            sources.append({
                "id": f"custom_{cs['id']}",
                "name": cs["name"],
                "url": feed_url,
                "category": cs["category"],
                "region": cs["region"],
                "relevance": cs["relevance"],
                "language": cs["language"],
                "paid": False,
                "is_custom": True,
                "custom_id": cs["id"]
            })
    except Exception:
        pass
    return sources

# ── Datenbank ─────────────────────────────────────────────────────────────────
def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()

    c.execute("""CREATE TABLE IF NOT EXISTS articles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        summary TEXT,
        url TEXT UNIQUE,
        source TEXT,
        source_id TEXT,
        category TEXT,
        region TEXT DEFAULT '',
        relevance INTEGER DEFAULT 5,
        fetched_date TEXT,
        is_favorite INTEGER DEFAULT 0,
        language TEXT DEFAULT 'en'
    )""")

    # ── Langzeit-Gedächtnis ────────────────────────────────────────────────────
    # Entitäten/Themen die aus Artikeln extrahiert werden
    c.execute("""CREATE TABLE IF NOT EXISTS entity_tags (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        article_id INTEGER,
        entity TEXT,
        entity_type TEXT DEFAULT 'topic',
        FOREIGN KEY(article_id) REFERENCES articles(id)
    )""")

    # Tägliche Häufigkeit jeder Entität
    c.execute("""CREATE TABLE IF NOT EXISTS daily_trends (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        date TEXT,
        entity TEXT,
        count INTEGER DEFAULT 1,
        avg_relevance REAL DEFAULT 5,
        sources TEXT DEFAULT '',
        UNIQUE(date, entity)
    )""")

    # Gespeicherte Prognosen
    c.execute("""CREATE TABLE IF NOT EXISTS prognoses (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        topic TEXT,
        timeframe TEXT,
        language TEXT DEFAULT 'de',
        headline TEXT,
        current_situation TEXT,
        driving_forces TEXT,
        prognosis_3m TEXT,
        prognosis_6m TEXT,
        prognosis_12m TEXT,
        autodoc_impact TEXT,
        confidence INTEGER DEFAULT 5,
        connected_articles TEXT,
        created_date TEXT
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS generated_articles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        headline TEXT,
        subheadline TEXT,
        meta_description TEXT,
        body TEXT,
        sources_used TEXT,
        created_date TEXT,
        language TEXT DEFAULT 'en'
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_topics (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        keywords TEXT,
        color TEXT DEFAULT '#F85A00',
        relevance_min INTEGER DEFAULT 1,
        active INTEGER DEFAULT 1,
        sort_order INTEGER DEFAULT 0,
        name_en TEXT DEFAULT '',
        name_uk TEXT DEFAULT '',
        name_ru TEXT DEFAULT ''
    )""")
    # Migration: Spalten hinzufügen falls Tabelle bereits existiert
    for col in ['name_en', 'name_uk', 'name_ru']:
        try:
            c.execute(f"ALTER TABLE user_topics ADD COLUMN {col} TEXT DEFAULT ''")
        except Exception:
            pass

    # Duplikate bereinigen: nur den ersten Eintrag pro Name behalten
    c.execute("""
        DELETE FROM user_topics WHERE id NOT IN (
            SELECT MIN(id) FROM user_topics GROUP BY name
        )
    """)
    # Doppelte eigene Quellen bereinigen (gleicher Name → neuesten behalten)
    c.execute("""
        DELETE FROM custom_sources WHERE id NOT IN (
            SELECT MAX(id) FROM custom_sources GROUP BY name
        )
    """)

    # Eingebaute Übersetzungen für Standard-Themen setzen (falls noch leer)
    builtin_translations = [
        ('Gesetzgebung & Recht',  'Legislation & Law',       'Законодавство та право',  'Законодательство и право'),
        ('Technische Normen',     'Technical Standards',     'Технічні норми',          'Технические нормы'),
        ('Marktdaten & Zahlen',   'Market Data & Figures',   'Ринкові дані та цифри',   'Рыночные данные и цифры'),
        ('Branchennews',          'Industry News',           'Новини галузі',           'Новости отрасли'),
        ('Elektromobilität',      'Electric Mobility',       'Електромобільність',      'Электромобильность'),
        ('Lieferkette',           'Supply Chain',            'Ланцюг поставок',         'Цепочка поставок'),
        ('Rohstoffe & Preise',    'Raw Materials & Prices',  'Сировина та ціни',        'Сырьё и цены'),
        ('Wettbewerber',          'Competitors',             'Конкуренти',              'Конкуренты'),
        ('Digital & SEO',         'Digital & SEO',           'Digital & SEO',           'Digital & SEO'),
        ('Nachhaltigkeit',        'Sustainability',          'Сталий розвиток',         'Устойчивость'),
    ]
    for name, en, uk, ru in builtin_translations:
        c.execute("""
            UPDATE user_topics SET name_en=?, name_uk=?, name_ru=?
            WHERE name=? AND (name_en='' OR name_en IS NULL)
        """, (en, uk, ru, name))

    c.execute("""CREATE TABLE IF NOT EXISTS source_credentials (
        source_id TEXT PRIMARY KEY,
        username TEXT DEFAULT '',
        password_b64 TEXT DEFAULT '',
        api_key TEXT DEFAULT '',
        subscriber_rss TEXT DEFAULT '',
        cookie TEXT DEFAULT '',
        active INTEGER DEFAULT 0
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS custom_sources (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        url TEXT NOT NULL,
        category TEXT DEFAULT 'Branchennews',
        region TEXT DEFAULT 'GLOBAL',
        relevance INTEGER DEFAULT 5,
        language TEXT DEFAULT 'en',
        active INTEGER DEFAULT 1,
        added_date TEXT,
        subscriber_rss TEXT DEFAULT '',
        api_key TEXT DEFAULT '',
        username TEXT DEFAULT '',
        password_b64 TEXT DEFAULT '',
        cookie TEXT DEFAULT ''
    )""")
    # Migration: Auth-Spalten hinzufügen falls noch nicht vorhanden
    for col in ['subscriber_rss', 'api_key', 'username', 'password_b64', 'cookie']:
        try:
            c.execute(f"ALTER TABLE custom_sources ADD COLUMN {col} TEXT DEFAULT ''")
        except Exception:
            pass

    c.execute("""CREATE TABLE IF NOT EXISTS app_settings (
        key TEXT PRIMARY KEY,
        value TEXT
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS category_settings (
        category TEXT PRIMARY KEY,
        priority INTEGER DEFAULT 5,
        active INTEGER DEFAULT 1
    )""")

    # Standard-Einstellungen
    defaults = {
        "language": "de",
        "relevance_min": "1",
        "refresh_hour": "7",
        "articles_per_feed": "10"
    }
    for k, v in defaults.items():
        c.execute("INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)", (k, v))

    # Standard-Kategorie-Prioritäten (können vom User angepasst werden)
    default_cat_priorities = [
        ("Gesetzgebung", 7), ("Technische Normen", 5), ("Marktdaten", 6),
        ("Branchennews", 7), ("Wirtschaft", 5), ("Elektromobilität", 9),
        ("Lieferkette", 6), ("Rohstoffe", 5), ("Aftermarket", 8),
        ("Digital & SEO", 8), ("Nachhaltigkeit", 6), ("Technologie", 7),
        ("Gesetzgebung & Recht", 7), ("Marktdaten & Zahlen", 6),
        ("Rohstoffe & Preise", 5), ("Digital & SEO", 8),
    ]
    for cat, prio in default_cat_priorities:
        c.execute("INSERT OR IGNORE INTO category_settings (category, priority) VALUES (?, ?)", (cat, prio))

    # Standard-Themen (AUTODOC-relevant)
    default_topics = [
        ("Gesetzgebung & Recht", "gesetz,recht,verordnung,regulation,law,legal,bundestag,EU-Recht", "#2196F3", 1),
        ("Technische Normen", "norm,standard,technisch,DIN,ISO,ECE,typgenehmigung", "#9C27B0", 1),
        ("Marktdaten & Zahlen", "markt,zulassung,statistik,sales,market,data,zahlen", "#4CAF50", 1),
        ("Branchennews", "automotive,auto,fahrzeug,kfz,industry,news", "#FF9800", 1),
        ("Elektromobilität", "elektro,EV,batterie,electric,hybrid,charging,laden", "#00BCD4", 1),
        ("Lieferkette", "lieferkette,supply chain,logistik,logistics,verfügbarkeit", "#795548", 1),
        ("Rohstoffe & Preise", "stahl,aluminium,rohstoff,steel,raw material,preis,price", "#FF5722", 1),
        ("Wettbewerber", "wettbewerber,competitor,LKQ,Mekonomen,ATU,autodistribution", "#607D8B", 1),
        ("Digital & SEO", "SEO,digital,ecommerce,google,shop,marketplace,amazon", "#009688", 1),
        ("Nachhaltigkeit", "nachhaltig,CO2,emission,green,sustainability,environment", "#8BC34A", 1),
    ]
    for name, keywords, color, active in default_topics:
        c.execute("INSERT OR IGNORE INTO user_topics (name, keywords, color, active) VALUES (?, ?, ?, ?)",
                  (name, keywords, color, active))

    conn.commit()
    conn.close()

def get_db():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")  # Mehrere gleichzeitige Leser erlauben
    return conn

def get_setting(key, default=""):
    conn = get_db()
    row = conn.execute("SELECT value FROM app_settings WHERE key = ?", (key,)).fetchone()
    conn.close()
    return row["value"] if row else default

def set_setting(key, value):
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO app_settings (key, value) VALUES (?, ?)", (key, str(value)))
    conn.commit()
    conn.close()

# ── Credential-Hilfsfunktionen ────────────────────────────────────────────────
def encode_password(pw):
    return base64.b64encode(pw.encode()).decode() if pw else ""

def decode_password(pw_b64):
    try:
        return base64.b64decode(pw_b64.encode()).decode() if pw_b64 else ""
    except Exception:
        return ""

# ── Feed-Abruf: kostenlos ─────────────────────────────────────────────────────
def fetch_free_source(source, max_items=10):
    try:
        feed = feedparser.parse(source["url"], request_headers={
            "User-Agent": "Mozilla/5.0 (compatible; AutodocInfoAgent/2.0)"
        })
        items = []
        for entry in feed.entries[:max_items]:
            title = entry.get("title", "").strip()
            url = entry.get("link", "")
            summary = entry.get("summary", entry.get("description", ""))
            summary = re.sub(r"<[^>]+>", "", summary).strip()[:600]
            if title and url:
                items.append({
                    "title": title, "url": url, "summary": summary,
                    "source": source["name"], "source_id": source.get("id", ""),
                    "category": source.get("category", ""), "region": source.get("region", ""),
                    "relevance": source.get("relevance", 5), "language": source.get("language", "en")
                })
        return items
    except Exception as e:
        print(f"  ✗ {source['name']}: {e}")
        return []

# ── Feed-Abruf: bezahlt (subscriber-RSS) ─────────────────────────────────────
def fetch_paid_source(source, creds):
    """Versucht Bezahlquelle via subscriber-RSS oder Session-Login zu erreichen."""
    # Methode 1: Subscriber-RSS URL (bevorzugt)
    if creds.get("subscriber_rss"):
        fake_src = dict(source)
        fake_src["url"] = creds["subscriber_rss"]
        return fetch_free_source(fake_src)

    # Methode 2: Cookie-basierter Zugriff
    if creds.get("cookie"):
        try:
            fake_src = dict(source)
            feed = feedparser.parse(source["url"], request_headers={
                "User-Agent": "Mozilla/5.0 (compatible; AutodocInfoAgent/2.0)",
                "Cookie": creds["cookie"]
            })
            items = []
            for entry in feed.entries[:10]:
                title = entry.get("title", "").strip()
                url = entry.get("link", "")
                summary = re.sub(r"<[^>]+>", "", entry.get("summary", "")).strip()[:600]
                if title and url:
                    items.append({
                        "title": title, "url": url, "summary": summary,
                        "source": source["name"], "source_id": source.get("id", ""),
                        "category": source.get("category", ""), "region": source.get("region", ""),
                        "relevance": source.get("relevance", 5), "language": source.get("language", "en")
                    })
            return items
        except Exception:
            pass

    # Methode 3: Benutzername/Passwort Session-Login
    if creds.get("username") and creds.get("password_b64"):
        try:
            session = requests.Session()
            session.headers.update({
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            })
            password = decode_password(creds["password_b64"])
            login_url = source.get("login_url", "")
            if login_url:
                session.post(login_url, data={
                    "username": creds["username"], "email": creds["username"],
                    "password": password, "login": "1"
                }, timeout=10)
                # Versuche RSS nach Login
                resp = session.get(source["url"], timeout=10)
                feed = feedparser.parse(resp.text)
                items = []
                for entry in feed.entries[:10]:
                    title = entry.get("title", "").strip()
                    url = entry.get("link", "")
                    summary = re.sub(r"<[^>]+>", "", entry.get("summary", "")).strip()[:600]
                    if title and url:
                        items.append({
                            "title": title, "url": url, "summary": summary,
                            "source": source["name"], "source_id": source.get("id", ""),
                            "category": source.get("category", ""), "region": source.get("region", ""),
                            "relevance": source.get("relevance", 5), "language": source.get("language", "en")
                        })
                if items:
                    return items
        except Exception as e:
            print(f"  ✗ Login fehlgeschlagen für {source['name']}: {e}")

    return []

# ── Hauptfunktion: alle Feeds holen (parallel) ───────────────────────────────
def fetch_feeds():
    global feed_status
    feed_status["running"] = True
    feed_status["new_count"] = 0
    print(f"\n[{datetime.now().strftime('%H:%M:%S')}] ── Feed-Abruf gestartet (parallel) ──")

    sources = load_sources()
    config = load_config()
    max_items = int(get_setting("articles_per_feed", "10"))
    today = date.today().isoformat()

    # Credentials laden (kostenpflichtige Standard-Quellen)
    conn = get_db()
    all_creds = {}
    for row in conn.execute("SELECT * FROM source_credentials WHERE active = 1").fetchall():
        all_creds[row["source_id"]] = dict(row)
    # Credentials für eigene Quellen aus custom_sources Tabelle laden
    custom_creds = {}
    for row in conn.execute("SELECT * FROM custom_sources WHERE active = 1").fetchall():
        cs = dict(row)
        cid = f"custom_{cs['id']}"
        custom_creds[cid] = {
            "cookie": cs.get("cookie", ""),
            "username": cs.get("username", ""),
            "api_key": cs.get("api_key", ""),
            "subscriber_rss": cs.get("subscriber_rss", ""),
            "password_b64": cs.get("password_b64", ""),
            "active": 1
        }
    conn.close()

    # Hilfsfunktion für eine Quelle
    def fetch_one(source):
        source_id = source.get("id", "")
        is_paid = source.get("paid", False)
        is_custom = source.get("is_custom", False)
        if is_paid:
            creds = all_creds.get(source_id, {})
            if not creds:
                return []
            return fetch_paid_source(source, creds)
        elif is_custom:
            # Eigene Quelle: cookie-basiert wenn vorhanden, sonst normaler Fetch
            creds = custom_creds.get(source_id, {})
            if creds.get("cookie"):
                return fetch_paid_source(source, creds)
            return fetch_free_source(source, max_items)
        else:
            return fetch_free_source(source, max_items)

    # Parallel fetchen mit 20 Threads gleichzeitig
    all_items = []
    with ThreadPoolExecutor(max_workers=20) as executor:
        futures = {executor.submit(fetch_one, src): src for src in sources}
        for future in as_completed(futures):
            try:
                items = future.result(timeout=15)
                all_items.extend(items)
            except Exception:
                pass

    # Alle Artikel in DB speichern
    conn = get_db()
    new_count = 0
    for item in all_items:
        try:
            conn.execute("""
                INSERT OR IGNORE INTO articles
                (title, summary, url, source, source_id, category, region, relevance, fetched_date, language)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (item["title"], item["summary"], item["url"], item["source"],
                  item["source_id"], item["category"], item["region"],
                  item["relevance"], today, item["language"]))
            new_count += 1
        except Exception:
            pass
    conn.commit()
    conn.close()

    feed_status["new_count"] = new_count
    feed_status["last_fetch"] = datetime.now().strftime('%H:%M')
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {new_count} neue Artikel gespeichert.")

    # KI-Bewertung + Entity-Extraktion
    api_key = config.get("claude_api_key", "")
    if api_key and api_key != "DEIN_API_KEY_HIER":
        rate_articles_with_ai(api_key)
        extract_entities_and_trends(api_key)

    feed_status["running"] = False

# ── KI-Relevanz-Bewertung ────────────────────────────────────────────────────
def rate_articles_with_ai(api_key):
    today = date.today().isoformat()
    conn = get_db()
    articles = conn.execute("""
        SELECT id, title, summary, category FROM articles
        WHERE fetched_date = ? AND relevance = 5
        LIMIT 40
    """, (today,)).fetchall()

    if not articles:
        conn.close()
        return

    # Themen für Kontext laden
    topics = conn.execute("SELECT name, keywords FROM user_topics WHERE active = 1").fetchall()
    topics_text = ", ".join([f"{t['name']} ({t['keywords'][:50]})" for t in topics])

    client = anthropic.Anthropic(api_key=api_key)
    articles_text = "\n".join([
        f"ID {a['id']}: [{a['category']}] {a['title']} — {a['summary'][:120]}"
        for a in articles
    ])

    try:
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=600,
            messages=[{"role": "user", "content":
                f"""Rate these articles by relevance for AUTODOC SE
(European online retailer for automotive spare parts, 26 markets).
User topics of interest: {topics_text}
Scale 1-10 (10=highest relevance).
Reply ONLY with JSON: {{"ratings": [{{"id": ID, "relevance": NUMBER}}, ...]}}
Articles:\n{articles_text}"""}]
        )
        result = json.loads(message.content[0].text)
        for r in result.get("ratings", []):
            conn.execute("UPDATE articles SET relevance = ? WHERE id = ?",
                        (r["relevance"], r["id"]))
        conn.commit()
        print(f"  KI: {len(result.get('ratings', []))} Artikel bewertet.")
    except Exception as e:
        print(f"  KI-Bewertung Fehler: {e}")
    conn.close()

# ── Flask App ─────────────────────────────────────────────────────────────────
app = Flask(__name__, template_folder="templates", static_folder="static")

@app.route("/")
def index():
    return render_template("index.html")

# ── Artikel ───────────────────────────────────────────────────────────────────
@app.route("/api/articles")
def api_articles():
    today = date.today().isoformat()
    category = request.args.get("category", "all")
    topic_id = request.args.get("topic_id")
    region = request.args.get("region", "all")
    relevance_min = int(request.args.get("relevance_min", get_setting("relevance_min", "1")))
    conn = get_db()

    # Wenn heute keine Artikel: neuestes Datum als Fallback
    count_today = conn.execute(
        "SELECT COUNT(*) FROM articles WHERE fetched_date = ?", (today,)).fetchone()[0]
    if count_today == 0:
        latest = conn.execute("SELECT MAX(fetched_date) FROM articles").fetchone()[0]
        display_date = latest if latest else today
    else:
        display_date = today

    # Sortierung: Effektive Relevanz = 60% KI-Bewertung + 40% Kategorie-Priorität des Users
    # Beispiel: KI=6, Kategorie-Priorität=9 → effektiv = 0.6*6 + 0.4*9 = 3.6+3.6 = 7.2
    query = """SELECT a.*,
        min(10.0, a.relevance * 0.6 + COALESCE(cs.priority, 5) * 0.4) as eff_rel
        FROM articles a
        LEFT JOIN category_settings cs ON a.category = cs.category
        WHERE a.fetched_date = ? AND a.relevance >= ?"""
    params = [display_date, relevance_min]

    if category != "all":
        query += " AND a.category = ?"
        params.append(category)
    if region == "REGIONAL":
        # Alle nicht-globalen Regionen (DE, EU, UK, FR, IT, ES, PL, UA, RU usw.)
        query += " AND a.region != 'GLOBAL'"
    elif region != "all":
        query += " AND a.region = ?"
        params.append(region)
    if topic_id:
        topic = conn.execute("SELECT keywords FROM user_topics WHERE id = ?",
                             (topic_id,)).fetchone()
        if topic:
            kws = [k.strip() for k in topic["keywords"].split(",") if k.strip()]
            if kws:
                kw_conditions = " OR ".join(["(a.title LIKE ? OR a.summary LIKE ? OR a.category LIKE ?)"] * len(kws))
                query += f" AND ({kw_conditions})"
                for kw in kws:
                    params += [f"%{kw}%", f"%{kw}%", f"%{kw}%"]

    query += " ORDER BY eff_rel DESC, a.id DESC LIMIT 100"
    articles = conn.execute(query, params).fetchall()
    conn.close()
    return jsonify([dict(a) for a in articles])

@app.route("/api/favorites")
def api_favorites():
    conn = get_db()
    articles = conn.execute("""
        SELECT * FROM articles WHERE is_favorite = 1
        ORDER BY fetched_date DESC, relevance DESC
    """).fetchall()
    conn.close()
    return jsonify([dict(a) for a in articles])

@app.route("/api/toggle_favorite", methods=["POST"])
def toggle_favorite():
    article_id = request.json.get("id")
    conn = get_db()
    cur = conn.execute("SELECT is_favorite FROM articles WHERE id = ?", (article_id,)).fetchone()
    if cur:
        conn.execute("UPDATE articles SET is_favorite = ? WHERE id = ?",
                    (0 if cur["is_favorite"] else 1, article_id))
        conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/refresh", methods=["POST"])
def api_refresh():
    threading.Thread(target=fetch_feeds, daemon=True).start()
    return jsonify({"status": "Feeds werden abgerufen..."})

# ── Themen ────────────────────────────────────────────────────────────────────
@app.route("/api/topics", methods=["GET"])
def get_topics():
    conn = get_db()
    topics = conn.execute("SELECT * FROM user_topics ORDER BY sort_order, id").fetchall()
    conn.close()
    return jsonify([dict(t) for t in topics])

@app.route("/api/topics/dedup", methods=["POST"])
def dedup_topics():
    """Einmalige Bereinigung von Duplikaten — wird beim App-Start automatisch aufgerufen."""
    conn = get_db()
    deleted = conn.execute("""
        DELETE FROM user_topics WHERE id NOT IN (
            SELECT MIN(id) FROM user_topics GROUP BY name
        )
    """).rowcount
    conn.commit()
    conn.close()
    return jsonify({"deleted": deleted})

@app.route("/api/topics", methods=["POST"])
def create_topic():
    conn = get_db()
    data = request.json
    new_name = data.get("name", "Neues Thema").strip()
    # Doppelten Namen verhindern
    existing = conn.execute(
        "SELECT id FROM user_topics WHERE LOWER(name) = LOWER(?)", (new_name,)
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": f"Thema '{new_name}' existiert bereits."}), 400
    count = conn.execute("SELECT COUNT(*) FROM user_topics").fetchone()[0]
    if count >= 50:
        conn.close()
        return jsonify({"error": "Maximum 50 Themen erreicht."}), 400
    conn.execute("""
        INSERT INTO user_topics (name, keywords, color, relevance_min, active, sort_order, name_en, name_uk, name_ru)
        VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?)
    """, (new_name, data.get("keywords", ""),
          data.get("color", "#F85A00"), data.get("relevance_min", 1), count,
          data.get("name_en", ""), data.get("name_uk", ""), data.get("name_ru", "")))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/topics/<int:topic_id>", methods=["PUT"])
def update_topic(topic_id):
    data = request.json
    conn = get_db()
    conn.execute("""
        UPDATE user_topics SET name=?, keywords=?, color=?, relevance_min=?, active=?,
            name_en=?, name_uk=?, name_ru=?
        WHERE id=?
    """, (data.get("name"), data.get("keywords"), data.get("color"),
          data.get("relevance_min", 1), data.get("active", 1),
          data.get("name_en", ""), data.get("name_uk", ""), data.get("name_ru", ""),
          topic_id))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/topics/<int:topic_id>", methods=["DELETE"])
def delete_topic(topic_id):
    conn = get_db()
    conn.execute("DELETE FROM user_topics WHERE id = ?", (topic_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# ── Einstellungen ─────────────────────────────────────────────────────────────
@app.route("/api/settings", methods=["GET"])
def get_settings():
    conn = get_db()
    rows = conn.execute("SELECT key, value FROM app_settings").fetchall()
    conn.close()
    settings = {r["key"]: r["value"] for r in rows}
    config = load_config()
    settings["claude_api_key"] = config.get("claude_api_key", "")
    return jsonify(settings)

@app.route("/api/settings", methods=["POST"])
def save_settings():
    data = request.json
    for key in ["language", "relevance_min", "refresh_hour", "articles_per_feed"]:
        if key in data:
            set_setting(key, data[key])
    if "claude_api_key" in data:
        config = load_config()
        config["claude_api_key"] = data["claude_api_key"]
        save_config(config)
    return jsonify({"success": True})

# ── Quellen & Credentials ─────────────────────────────────────────────────────
@app.route("/api/sources", methods=["GET"])
def get_sources():
    try:
        sources = load_sources()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    conn = get_db()
    creds_rows = conn.execute("SELECT * FROM source_credentials").fetchall()
    conn.close()
    creds_map = {r["source_id"]: dict(r) for r in creds_rows}
    for s in sources:
        sid = s.get("id", "")
        cred = creds_map.get(sid, {})
        s["has_credentials"] = bool(cred.get("active"))
        s["subscriber_rss"] = cred.get("subscriber_rss", "")
    return jsonify(sources)

@app.route("/api/credentials/<source_id>", methods=["GET"])
def get_credentials(source_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM source_credentials WHERE source_id = ?",
                       (source_id,)).fetchone()
    conn.close()
    if row:
        d = dict(row)
        d["password_b64"] = ""  # Passwort nicht zurückgeben
        return jsonify(d)
    return jsonify({"source_id": source_id, "username": "", "api_key": "",
                    "subscriber_rss": "", "cookie": "", "active": 0})

@app.route("/api/credentials/<source_id>", methods=["POST"])
def save_credentials(source_id):
    data = request.json
    password = data.get("password", "")
    conn = get_db()
    existing = conn.execute("SELECT password_b64 FROM source_credentials WHERE source_id = ?",
                            (source_id,)).fetchone()
    # Passwort nur updaten wenn neu eingegeben
    if password:
        pw_b64 = encode_password(password)
    elif existing:
        pw_b64 = existing["password_b64"]
    else:
        pw_b64 = ""

    conn.execute("""
        INSERT OR REPLACE INTO source_credentials
        (source_id, username, password_b64, api_key, subscriber_rss, cookie, active)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (source_id, data.get("username", ""), pw_b64,
          data.get("api_key", ""), data.get("subscriber_rss", ""),
          data.get("cookie", ""), 1 if data.get("active") else 0))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# ── Entity-Extraktion & Trend-Tracking ────────────────────────────────────────
def extract_entities_and_trends(api_key: str):
    """Extrahiert Entitäten aus neuen Artikeln und trackt tägliche Trends."""
    today = date.today().isoformat()
    conn = get_db()

    # Neue Artikel ohne Entity-Tags
    articles = conn.execute("""
        SELECT a.id, a.title, a.summary, a.category, a.source, a.relevance
        FROM articles a
        LEFT JOIN entity_tags e ON a.id = e.article_id
        WHERE a.fetched_date = ? AND e.id IS NULL
        LIMIT 50
    """, (today,)).fetchall()

    if not articles:
        conn.close()
        return

    client = anthropic.Anthropic(api_key=api_key)
    articles_text = "\n".join([
        f"ID {a['id']}: {a['title']} — {a['summary'][:150]}"
        for a in articles
    ])

    try:
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1500,
            messages=[{"role": "user", "content":
                f"""Extract key entities (topics, companies, technologies, regulations, locations, events)
from these automotive/business news articles. Focus on entities relevant to AUTODOC SE
(European automotive spare parts retailer).

Return ONLY valid JSON:
{{
  "articles": [
    {{
      "id": <article_id>,
      "entities": [
        {{"name": "entity name", "type": "topic|company|regulation|technology|location|event"}},
        ...
      ]
    }},
    ...
  ]
}}

Articles:
{articles_text}"""}]
        )

        result = json.loads(message.content[0].text)

        for art_data in result.get("articles", []):
            art_id = art_data["id"]
            for entity in art_data.get("entities", []):
                name = entity["name"].strip()
                etype = entity.get("type", "topic")
                if len(name) > 2:
                    try:
                        conn.execute(
                            "INSERT OR IGNORE INTO entity_tags (article_id, entity, entity_type) VALUES (?, ?, ?)",
                            (art_id, name, etype)
                        )
                    except Exception:
                        pass

        conn.commit()
        print(f"  Entities: {sum(len(a.get('entities',[])) for a in result.get('articles',[]))} extrahiert.")

    except Exception as e:
        print(f"  Entity-Extraktion Fehler: {e}")

    # Tägliche Trends aktualisieren
    try:
        conn.execute("""
            INSERT INTO daily_trends (date, entity, count, avg_relevance, sources)
            SELECT
                a.fetched_date,
                e.entity,
                COUNT(*) as count,
                AVG(a.relevance) as avg_rel,
                GROUP_CONCAT(DISTINCT a.source)
            FROM entity_tags e
            JOIN articles a ON e.article_id = a.id
            WHERE a.fetched_date = ?
            GROUP BY a.fetched_date, e.entity
            ON CONFLICT(date, entity) DO UPDATE SET
                count = excluded.count,
                avg_relevance = excluded.avg_relevance,
                sources = excluded.sources
        """, (today,))
        conn.commit()
    except Exception as e:
        print(f"  Trend-Update Fehler: {e}")

    conn.close()


# ── Bild-Hilfsfunktionen ──────────────────────────────────────────────────────
# Kuratierte Unsplash-Foto-IDs nach Thema (kein API-Key nötig)
CURATED_IMAGES = {
    "automotive":    ["photo-1492144534655-ae79c964c9d7", "photo-1568605117036-5fe5e7bab0b7",
                      "photo-1503376780353-7e6692767b70", "photo-1511919884226-fd3cad34687c"],
    "car_parts":     ["photo-1558618666-fcd25c85cd64", "photo-1486262715619-67b85e0b08d3",
                      "photo-1635070041078-e363dbe005cb", "photo-1609521263047-f8f205293f24"],
    "ev_electric":   ["photo-1593941707874-ef25b8b4a92b", "photo-1617788138017-80ad40651399",
                      "photo-1647166545149-8dc68ab6cafd", "photo-1619767886558-efdc259cde1a"],
    "warehouse":     ["photo-1586528116311-ad8dd3c8310d", "photo-1553413077-190dd305871c",
                      "photo-1578575437130-527eed3abbec", "photo-1587293852726-70cdb56c2866"],
    "market_data":   ["photo-1611974789855-9c2a0a7236a3", "photo-1543286386-2e659306cd6c",
                      "photo-1590283603385-17ffb3a7f29f", "photo-1460925895917-afdab827c52f"],
    "regulations":   ["photo-1450101499163-c8848c66ca85", "photo-1589829545856-d10d557cf95f",
                      "photo-1507003211169-0a1dd7228f2d", "photo-1521791136064-7986c2920216"],
    "supply_chain":  ["photo-1586528116162-c643bfec6887", "photo-1601584115197-04ecc0da31d7",
                      "photo-1494412519320-aa613dfb7738", "photo-1519003722824-194d4455a60c"],
    "raw_materials": ["photo-1504917595217-d4dc5ebe6122", "photo-1567789884554-0b844b597180",
                      "photo-1518443855757-dfadac7101ae", "photo-1605792657660-596af9009e82"],
    "digital":       ["photo-1460925895917-afdab827c52f", "photo-1518770660439-4636190af475",
                      "photo-1519389950473-47ba0277781c", "photo-1498050108023-c5249f4df085"],
    "sustainability":["photo-1473341304170-971dccb5ac1e", "photo-1497436072909-60f360e1d4b1",
                      "photo-1542601906990-b4d3fb778b09", "photo-1466611653911-95081537e5b7"],
    "workshop":      ["photo-1580274455191-1c62238fa333", "photo-1530046339160-ce3e530c7d2f",
                      "photo-1504222490345-c075b7011089", "photo-1562259929-b4e1fd3aef09"],
    "highway":       ["photo-1485291571150-772bcfc10da5", "photo-1449965408869-eaa3f722e40d",
                      "photo-1545558014-8692077e9b5c", "photo-1498036882173-b41c28a8ba34"],
    "business":      ["photo-1507003211169-0a1dd7228f2d", "photo-1552664730-d307ca884978",
                      "photo-1521791136064-7986c2920216", "photo-1454165804606-c3d57bc86b40"],
    "europe":        ["photo-1467269204594-9661b134dd2b", "photo-1541849546-216549ae216d",
                      "photo-1519677100203-a0e668c92439", "photo-1491557345352-5929e343eb89"],
    "technology":    ["photo-1518770660439-4636190af475", "photo-1485827404703-89b55fcc595e",
                      "photo-1526374965328-7f61d4dc18c5", "photo-1531297484001-80022131f5a1"],
    "engine":        ["photo-1486262715619-67b85e0b08d3", "photo-1609521263047-f8f205293f24",
                      "photo-1635070041078-e363dbe005cb", "photo-1558618666-fcd25c85cd64"],
    "tires":         ["photo-1591768793355-74d04bb6608f", "photo-1604999565976-8913ad2ddb7c",
                      "photo-1568772585407-9361f9bf3a87", "photo-1580273916550-e323be2ae537"],
    "trucks":        ["photo-1601584115197-04ecc0da31d7", "photo-1519003722824-194d4455a60c",
                      "photo-1494412519320-aa613dfb7738", "photo-1519519166026-64d4b9b16f94"],
    "factory":       ["photo-1565043589221-1a6fd9ae45c7", "photo-1504328345606-18bbc8c9d7d1",
                      "photo-1581091226825-a6a2a5aee158", "photo-1537462715879-360eeb61a0ad"],
    "ecommerce":     ["photo-1556742049-0cfed4f6a45d", "photo-1563013544-824ae1b704d3",
                      "photo-1607082348824-0a96f2a4b9da", "photo-1472851294608-062f824d29cc"],
}

def get_image_url(query_hint: str, width=1200, height=480, used_ids: set = None) -> str:
    """Gibt eine Unsplash-Bild-URL zurück. Vermeidet bereits verwendete Bilder."""
    if used_ids is None:
        used_ids = set()
    mapping = {
        "gesetzgebung": "regulations", "regulation": "regulations", "law": "regulations",
        "technisch": "engine", "technical": "engine", "norm": "regulations",
        "marktdaten": "market_data", "market": "market_data", "statistik": "market_data",
        "branche": "automotive", "news": "automotive", "automotive": "automotive",
        "elektro": "ev_electric", "electric": "ev_electric", "ev": "ev_electric",
        "lieferkette": "supply_chain", "supply": "supply_chain", "logistik": "warehouse",
        "rohstoff": "raw_materials", "material": "raw_materials", "stahl": "raw_materials",
        "digital": "digital", "seo": "digital", "ecommerce": "ecommerce",
        "nachhaltig": "sustainability", "green": "sustainability", "co2": "sustainability",
        "aftermarket": "workshop", "werkstatt": "workshop", "repair": "workshop",
        "technologie": "technology", "tech": "technology",
        "europa": "europe", "europe": "europe",
    }
    hint_lower = query_hint.lower()
    photo_key = "automotive"
    for keyword, key in mapping.items():
        if keyword in hint_lower:
            photo_key = key
            break
    pool = CURATED_IMAGES.get(photo_key, CURATED_IMAGES["automotive"])
    # Nicht bereits verwendete Bilder bevorzugen
    available = [p for p in pool if p not in used_ids]
    if not available:
        available = pool  # Alle verbraucht: Pool wiederverwenden
    photo_id = available[0]
    used_ids.add(photo_id)
    return f"https://images.unsplash.com/{photo_id}?w={width}&h={height}&fit=crop&q=80"


# ── Artikel-Erstellung ────────────────────────────────────────────────────────
LANG_INSTRUCTIONS = {
    "de": "Schreibe einen professionellen, publikationsreifen Artikel auf Deutsch",
    "en": "Write a professional, publication-ready article in English",
    "uk": "Напишіть професійну статтю, готову до публікації, українською мовою",
    "ru": "Напишите профессиональную статью, готовую к публикации, на русском языке",
}

@app.route("/api/generate_article", methods=["POST"])
def generate_article():
    config = load_config()
    api_key = config.get("claude_api_key", "")
    if not api_key or api_key == "DEIN_API_KEY_HIER":
        return jsonify({"error": "Kein Claude API-Key eingetragen."}), 400

    data = request.json
    selected_ids = data.get("article_ids", [])
    article_lang = data.get("article_language", "en")  # Artikelsprache unabhängig von UI
    today = date.today().isoformat()
    conn = get_db()

    if selected_ids:
        placeholders = ",".join("?" * len(selected_ids))
        articles = conn.execute(
            f"SELECT * FROM articles WHERE id IN ({placeholders})", selected_ids).fetchall()
    else:
        articles = conn.execute(
            "SELECT * FROM articles WHERE fetched_date = ? ORDER BY relevance DESC LIMIT 8",
            (today,)).fetchall()
        # Fallback: Wenn heute keine Artikel vorhanden, neuestes verfügbares Datum nehmen
        if not articles:
            latest = conn.execute(
                "SELECT MAX(fetched_date) FROM articles").fetchone()[0]
            if latest:
                articles = conn.execute(
                    "SELECT * FROM articles WHERE fetched_date = ? ORDER BY relevance DESC LIMIT 8",
                    (latest,)).fetchall()
                print(f"  Kein Artikel für heute ({today}), nutze Daten von {latest}")
    conn.close()

    if not articles:
        return jsonify({"error": "Keine Artikel in der Datenbank. Bitte zuerst Feeds aktualisieren."}), 404

    facts = "\n".join([
        f"- [{a['category']}] {a['title']}: {a['summary'][:250]}\n  Source: {a['source']} ({a['url']})"
        for a in articles
    ])
    main_category = articles[0]["category"] if articles else "automotive"
    lang_instruction = LANG_INSTRUCTIONS.get(article_lang, LANG_INSTRUCTIONS["en"])

    client = anthropic.Anthropic(api_key=api_key)
    prompt = f"""{lang_instruction} for AUTODOC SE (European e-commerce for automotive spare parts, 26 markets) based on the following news facts.

SOURCE FACTS:
{facts}

The article must be PUBLICATION-READY with this EXACT JSON structure — no extra text, only valid JSON:
{{
  "headline": "Compelling SEO headline (max 65 chars)",
  "subheadline": "Explanatory subheadline (max 130 chars)",
  "meta_description": "SEO meta description (max 155 chars)",
  "author_note": "Short editorial context (1 sentence, why this matters)",
  "hero_image_hint": "2-3 English keywords for hero image (e.g. 'electric vehicle charging')",
  "key_facts": [
    {{"icon": "📊", "label": "Fact label", "value": "Fact value"}},
    {{"icon": "🌍", "label": "Fact label", "value": "Fact value"}},
    {{"icon": "📈", "label": "Fact label", "value": "Fact value"}},
    {{"icon": "⚡", "label": "Fact label", "value": "Fact value"}}
  ],
  "intro": "Strong lead paragraph (3-4 sentences, captures the story)",
  "pull_quote": "One powerful, quotable sentence from the article",
  "sections": [
    {{
      "heading": "H2 section heading",
      "content": "2-3 paragraph text, detailed and informative",
      "image_hint": "2-3 English keywords for section image",
      "has_table": false
    }},
    {{
      "heading": "Data & Market Overview",
      "content": "Paragraph introducing the table below",
      "image_hint": "market data statistics",
      "has_table": true,
      "table": {{
        "caption": "Table caption",
        "headers": ["Column 1", "Column 2", "Column 3"],
        "rows": [
          ["Row 1 data", "Value", "Value"],
          ["Row 2 data", "Value", "Value"],
          ["Row 3 data", "Value", "Value"]
        ]
      }}
    }},
    {{
      "heading": "H2 section heading",
      "content": "2-3 paragraph text",
      "image_hint": "relevant keywords",
      "has_table": false
    }}
  ],
  "conclusion": "Strong closing paragraph with outlook (3-4 sentences)",
  "tags": ["tag1", "tag2", "tag3", "tag4", "tag5"],
  "sources": [
    {{"name": "Source Name", "url": "https://..."}}
  ]
}}

Rules:
- Write in {article_lang.upper()} throughout (headline, body, everything)
- Tables must contain real, plausible data derived from the facts
- Each section should have 2-3 substantial paragraphs
- Key facts must be specific numbers or short impactful statements
- Tone: Professional, factual, B2B-ready, no speculation"""

    def extract_json(text):
        """Extrahiert JSON robust aus KI-Antwort — auch aus Markdown-Codeblöcken."""
        t = text.strip()
        # Markdown-Codeblock entfernen: ```json ... ``` oder ``` ... ```
        t = re.sub(r'^```(?:json)?\s*', '', t, flags=re.IGNORECASE)
        t = re.sub(r'\s*```$', '', t)
        t = t.strip()
        # Direkt valides JSON?
        if t.startswith('{'):
            return t
        # JSON-Block aus gemischtem Text heraussuchen
        match = re.search(r'\{[\s\S]*\}', t)
        if match:
            return match.group(0)
        return None

    for attempt in range(1, 4):  # bis zu 3 Versuche
        try:
            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=8000,
                system="You are a JSON-only API. Always respond with a single valid JSON object. No markdown, no explanations, no code blocks — pure JSON only.",
                messages=[{"role": "user", "content": prompt}]
            )
            raw = message.content[0].text.strip()
            print(f"  Artikel-Antwort (Versuch {attempt}): {len(raw)} Zeichen, starts: {raw[:40]!r}")
            extracted = extract_json(raw)
            if not extracted:
                print(f"  Kein JSON gefunden (Versuch {attempt}), nochmal...")
                continue
            result = json.loads(extracted)
            break  # Erfolg
        except json.JSONDecodeError as e:
            print(f"  JSON-Fehler Versuch {attempt}: {e} | Anfang: {raw[:300]}")
            if attempt == 3:
                return jsonify({"error": "JSON-Fehler beim Parsen der KI-Antwort. Bitte nochmals versuchen."}), 500
            continue
        except Exception as e:
            print(f"  Artikel-Fehler: {type(e).__name__}: {e}")
            return jsonify({"error": f"{type(e).__name__}: {str(e)}"}), 500
    else:
        return jsonify({"error": "KI-Antwort enthält kein gültiges JSON nach 3 Versuchen."}), 500

    # Bild-URLs hinzufügen — shared used_ids verhindert Duplikate im Artikel
    used_img_ids = set()
    result["hero_image_url"] = get_image_url(
        result.get("hero_image_hint", main_category), used_ids=used_img_ids)
    for section in result.get("sections", []):
        section["image_url"] = get_image_url(
            section.get("image_hint", main_category), width=800, height=300,
            used_ids=used_img_ids)

    # In DB speichern
    db_conn = get_db()
    db_conn.execute("""
        INSERT INTO generated_articles
        (headline, subheadline, meta_description, body, sources_used, created_date, language)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (result.get("headline"), result.get("subheadline"), result.get("meta_description"),
          json.dumps(result), json.dumps(result.get("sources", [])), today, article_lang))
    last_id = db_conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    db_conn.commit()
    db_conn.close()
    result["id"] = last_id
    print(f"  ✓ Artikel erstellt: {result.get('headline','')[:60]}")
    return jsonify(result)


# ── Trend & Prognose Endpunkte ────────────────────────────────────────────────

@app.route("/api/trends")
def api_trends():
    """Gibt die wichtigsten Trends der letzten N Tage zurück."""
    days = int(request.args.get("days", 30))
    limit = int(request.args.get("limit", 50))
    conn = get_db()

    trends = conn.execute("""
        SELECT
            entity,
            SUM(count) as total_mentions,
            AVG(avg_relevance) as avg_rel,
            COUNT(DISTINCT date) as days_active,
            MIN(date) as first_seen,
            MAX(date) as last_seen,
            GROUP_CONCAT(DISTINCT sources) as all_sources
        FROM daily_trends
        WHERE date >= date('now', ? || ' days')
        GROUP BY entity
        HAVING total_mentions >= 2
        ORDER BY total_mentions DESC, avg_rel DESC
        LIMIT ?
    """, (f"-{days}", limit)).fetchall()
    conn.close()
    return jsonify([dict(t) for t in trends])


@app.route("/api/trend_timeline/<path:entity>")
def api_trend_timeline(entity):
    """Zeigt wie oft eine Entität über die Zeit erwähnt wurde."""
    days = int(request.args.get("days", 60))
    conn = get_db()

    timeline = conn.execute("""
        SELECT date, count, avg_relevance, sources
        FROM daily_trends
        WHERE entity = ? AND date >= date('now', ? || ' days')
        ORDER BY date ASC
    """, (entity, f"-{days}")).fetchall()

    # Zugehörige Artikel
    articles = conn.execute("""
        SELECT DISTINCT a.id, a.title, a.url, a.source, a.fetched_date, a.relevance, a.summary
        FROM articles a
        JOIN entity_tags e ON a.id = e.article_id
        WHERE e.entity = ? AND a.fetched_date >= date('now', ? || ' days')
        ORDER BY a.fetched_date DESC, a.relevance DESC
        LIMIT 30
    """, (entity, f"-{days}")).fetchall()

    conn.close()
    return jsonify({
        "entity": entity,
        "timeline": [dict(t) for t in timeline],
        "articles": [dict(a) for a in articles]
    })


@app.route("/api/prognosis", methods=["POST"])
def generate_prognosis():
    """Erstellt eine KI-gestützte Zukunftsprognose basierend auf historischen Artikeln."""
    config = load_config()
    api_key = config.get("claude_api_key", "")
    if not api_key or api_key == "DEIN_API_KEY_HIER":
        return jsonify({"error": "Kein Claude API-Key eingetragen."}), 400

    data = request.json
    topic = data.get("topic", "")
    days = int(data.get("days", 60))
    prog_lang = data.get("language", get_setting("language", "de"))

    if not topic:
        return jsonify({"error": "Kein Thema angegeben."}), 400

    conn = get_db()

    # Alle relevanten Artikel zu diesem Thema aus den letzten N Tagen sammeln
    # Suche nach Entity-Tags UND Keyword-Match in Titel/Summary
    articles_by_entity = conn.execute("""
        SELECT DISTINCT a.id, a.title, a.summary, a.source, a.fetched_date,
               a.relevance, a.category, a.region, a.url
        FROM articles a
        JOIN entity_tags e ON a.id = e.article_id
        WHERE LOWER(e.entity) LIKE LOWER(?)
          AND a.fetched_date >= date('now', ? || ' days')
        ORDER BY a.fetched_date DESC, a.relevance DESC
    """, (f"%{topic}%", f"-{days}")).fetchall()

    articles_by_keyword = conn.execute("""
        SELECT DISTINCT id, title, summary, source, fetched_date,
               relevance, category, region, url
        FROM articles
        WHERE (LOWER(title) LIKE LOWER(?) OR LOWER(summary) LIKE LOWER(?))
          AND fetched_date >= date('now', ? || ' days')
        ORDER BY fetched_date DESC, relevance DESC
        LIMIT 50
    """, (f"%{topic}%", f"%{topic}%", f"-{days}")).fetchall()

    # Zusammenführen und deduplizieren
    seen_ids = set()
    all_articles = []
    for a in list(articles_by_entity) + list(articles_by_keyword):
        if a["id"] not in seen_ids:
            seen_ids.add(a["id"])
            all_articles.append(dict(a))

    # Trend-Timeline für dieses Thema
    timeline = conn.execute("""
        SELECT date, SUM(count) as count, AVG(avg_relevance) as avg_rel
        FROM daily_trends
        WHERE LOWER(entity) LIKE LOWER(?)
          AND date >= date('now', ? || ' days')
        GROUP BY date
        ORDER BY date ASC
    """, (f"%{topic}%", f"-{days}")).fetchall()

    # Verwandte Entitäten finden (oft gemeinsam mit dem Thema erwähnt)
    related = conn.execute("""
        SELECT e2.entity, COUNT(*) as co_occurrences
        FROM entity_tags e1
        JOIN entity_tags e2 ON e1.article_id = e2.article_id
        WHERE LOWER(e1.entity) LIKE LOWER(?)
          AND LOWER(e2.entity) NOT LIKE LOWER(?)
        GROUP BY e2.entity
        ORDER BY co_occurrences DESC
        LIMIT 15
    """, (f"%{topic}%", f"%{topic}%")).fetchall()

    conn.close()

    if not all_articles:
        return jsonify({"error": f"Keine Artikel zu '{topic}' in den letzten {days} Tagen gefunden. Bitte zuerst Feeds aktualisieren."}), 404

    # Artikel-Timeline für den Prompt aufbereiten
    facts_by_date = {}
    for a in sorted(all_articles, key=lambda x: x["fetched_date"]):
        d = a["fetched_date"]
        if d not in facts_by_date:
            facts_by_date[d] = []
        facts_by_date[d].append(
            f"  [{a['source']} | REL:{a['relevance']}] {a['title']}: {a['summary'][:200]}"
        )

    timeline_text = "\n".join([
        f"{d}:\n" + "\n".join(items)
        for d, items in sorted(facts_by_date.items())
    ])

    related_text = ", ".join([f"{r['entity']} ({r['co_occurrences']}x)" for r in related])

    lang_instructions = {
        "de": "Antworte auf Deutsch",
        "en": "Reply in English",
        "uk": "Відповідай українською",
        "ru": "Отвечай на русском"
    }
    lang_instr = lang_instructions.get(prog_lang, lang_instructions["de"])

    client = anthropic.Anthropic(api_key=api_key)

    prompt = f"""Du bist ein Senior-Analyst für die Automobilbranche. {lang_instr}.

Analysiere die folgende Nachrichtenhistorie zum Thema "{topic}" aus den letzten {days} Tagen
und erstelle eine fundierte Zukunftsprognose für AUTODOC SE (europäischer Online-Händler
für Kfz-Ersatzteile, 26 Märkte, Hauptmärkte DE/PL/RU/UA/FR).

NACHRICHTENCHRONOLOGIE:
{timeline_text}

VERWANDTE THEMEN (häufig gemeinsam erwähnt): {related_text}

Erstelle eine tiefgehende Analyse. Verbinde scheinbar unzusammenhängende Ereignisse
(z.B. geopolitische Ereignisse → Rohstoffpreise → Fahrzeugbestand → Ersatzteilbedarf).
Denke in Ursache-Wirkung-Ketten.

Antworte NUR mit diesem JSON (kein anderer Text):
{{
  "headline": "Prägnante Überschrift der Prognose (max 80 Zeichen)",
  "confidence": <Konfidenz 1-10, 10=sehr sicher>,
  "current_situation": "Aktuelle Lage in 3-4 Sätzen. Was passiert gerade?",
  "key_connections": [
    {{
      "chain": "Ereignis A → Auswirkung B → Folge C",
      "explanation": "Warum diese Verbindung wichtig ist"
    }},
    {{
      "chain": "...",
      "explanation": "..."
    }}
  ],
  "driving_forces": [
    {{"force": "Treiber 1", "direction": "positiv|negativ|neutral", "strength": "stark|mittel|schwach"}},
    {{"force": "Treiber 2", "direction": "...", "strength": "..."}}
  ],
  "scenarios": {{
    "base": {{
      "label": "Basisszenario (wahrscheinlichstes)",
      "probability": "60-70%",
      "description": "Was passiert wenn alles normal weiterläuft"
    }},
    "bull": {{
      "label": "Positivszenario",
      "probability": "15-20%",
      "description": "Was passiert wenn günstige Faktoren eintreten"
    }},
    "bear": {{
      "label": "Negativszenario",
      "probability": "15-20%",
      "description": "Was passiert wenn ungünstige Faktoren eintreten"
    }}
  }},
  "prognosis_3m": "Prognose für die nächsten 3 Monate (konkret, mit Zahlen wo möglich)",
  "prognosis_6m": "Prognose für die nächsten 6 Monate",
  "prognosis_12m": "Prognose für die nächsten 12 Monate",
  "autodoc_impact": {{
    "summary": "Gesamtauswirkung auf AUTODOC in 2-3 Sätzen",
    "affected_categories": [
      {{"category": "Produktkategorie", "impact": "Beschreibung", "direction": "steigend|fallend|neutral"}}
    ],
    "actions": [
      "Empfohlene Maßnahme 1",
      "Empfohlene Maßnahme 2",
      "Empfohlene Maßnahme 3"
    ]
  }},
  "watch_signals": [
    "Signal das beobachtet werden sollte 1",
    "Signal 2",
    "Signal 3"
  ],
  "sources_used": {len(all_articles)}
}}"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = message.content[0].text.strip()
        if not raw.startswith('{'):
            match = re.search(r'\{[\s\S]*\}', raw)
            if match:
                raw = match.group(0)
        result = json.loads(raw)

        # In DB speichern
        db_conn = get_db()
        db_conn.execute("""
            INSERT INTO prognoses
            (topic, timeframe, language, headline, current_situation,
             driving_forces, prognosis_3m, prognosis_6m, prognosis_12m,
             autodoc_impact, confidence, connected_articles, created_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            topic, f"{days} Tage", prog_lang,
            result.get("headline", ""),
            result.get("current_situation", ""),
            json.dumps(result.get("driving_forces", [])),
            result.get("prognosis_3m", ""),
            result.get("prognosis_6m", ""),
            result.get("prognosis_12m", ""),
            json.dumps(result.get("autodoc_impact", {})),
            result.get("confidence", 5),
            json.dumps([a["id"] for a in all_articles]),
            date.today().isoformat()
        ))
        db_conn.commit()
        db_conn.close()

        result["articles_analyzed"] = len(all_articles)
        result["related_topics"] = [{"entity": r["entity"], "count": r["co_occurrences"]} for r in related]
        result["timeline_data"] = [dict(t) for t in timeline]
        return jsonify(result)

    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON-Fehler: {e}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/prognoses", methods=["GET"])
def list_prognoses():
    conn = get_db()
    rows = conn.execute("""
        SELECT id, topic, headline, confidence, language, created_date
        FROM prognoses ORDER BY id DESC LIMIT 20
    """).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/prognoses/<int:prog_id>", methods=["GET"])
def get_prognosis(prog_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM prognoses WHERE id = ?", (prog_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "Nicht gefunden"}), 404
    d = dict(row)
    d["driving_forces"] = json.loads(d.get("driving_forces") or "[]")
    d["autodoc_impact"] = json.loads(d.get("autodoc_impact") or "{}")
    d["connected_articles"] = json.loads(d.get("connected_articles") or "[]")
    return jsonify(d)


@app.route("/api/memory_stats")
def memory_stats():
    """Überblick über das gesamte gespeicherte Wissen."""
    conn = get_db()
    total_articles = conn.execute("SELECT COUNT(*) FROM articles").fetchone()[0]
    total_entities = conn.execute("SELECT COUNT(DISTINCT entity) FROM entity_tags").fetchone()[0]
    total_prognoses = conn.execute("SELECT COUNT(*) FROM prognoses").fetchone()[0]
    date_range = conn.execute(
        "SELECT MIN(fetched_date), MAX(fetched_date) FROM articles").fetchone()
    top_entities = conn.execute("""
        SELECT entity, COUNT(*) as c FROM entity_tags
        GROUP BY entity ORDER BY c DESC LIMIT 10
    """).fetchall()
    trend_count = conn.execute("SELECT COUNT(DISTINCT entity) FROM daily_trends").fetchone()[0]
    conn.close()
    return jsonify({
        "total_articles": total_articles,
        "total_entities": total_entities,
        "total_prognoses": total_prognoses,
        "trend_entities": trend_count,
        "date_from": date_range[0] or "—",
        "date_to": date_range[1] or "—",
        "top_entities": [dict(e) for e in top_entities]
    })


@app.route("/api/generated_articles", methods=["GET"])
def list_generated_articles():
    conn = get_db()
    rows = conn.execute("""
        SELECT id, headline, language, created_date FROM generated_articles
        ORDER BY id DESC LIMIT 20
    """).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/export_article/<int:article_id>", methods=["GET"])
def export_article(article_id):
    """Gibt einen HTML-Export des Artikels zurück."""
    conn = get_db()
    row = conn.execute("SELECT * FROM generated_articles WHERE id = ?", (article_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "Artikel nicht gefunden."}), 404

    data = json.loads(row["body"])
    html = render_article_html(data, row["language"])

    # Als Datei speichern
    filename = f"artikel_{article_id}_{row['language']}_{row['created_date']}.html"
    export_dir = os.path.join(BASE_DIR, "exports")
    os.makedirs(export_dir, exist_ok=True)
    filepath = os.path.join(export_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)

    from flask import send_file
    return send_file(filepath, as_attachment=True, download_name=filename,
                     mimetype="text/html")


@app.route("/api/export_article_docx/<int:article_id>", methods=["GET"])
def export_article_docx(article_id):
    """Word-Export (.docx) mit Bildern — generiert server-seitig."""
    import subprocess, sys
    try:
        from docx import Document
    except ImportError:
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "python-docx", "--quiet"],
                stderr=subprocess.DEVNULL
            )
            from docx import Document
        except Exception as e:
            return jsonify({"error": f"python-docx konnte nicht installiert werden: {e}"}), 500

    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO

    conn = get_db()
    row = conn.execute("SELECT * FROM generated_articles WHERE id = ?", (article_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "Artikel nicht gefunden."}), 404

    data = json.loads(row["body"])
    lang = row["language"]
    today_str = date.today().strftime("%d. %B %Y" if lang == "de" else "%B %d, %Y")

    ORANGE = RGBColor(0xF8, 0x5A, 0x00)
    DARK   = RGBColor(0x13, 0x1C, 0x20)
    MUTED  = RGBColor(0x66, 0x66, 0x66)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    PEACH  = RGBColor(0xFF, 0xBA, 0x96)

    def download_image(url):
        if not url:
            return None
        try:
            r = requests.get(url, timeout=12, headers={"User-Agent": "Mozilla/5.0"})
            if r.status_code == 200 and len(r.content) > 500:
                return BytesIO(r.content)
        except Exception:
            pass
        return None

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_para_border_left(para, hex_color="F85A00", sz="24", space="20"):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), sz)
        left.set(qn("w:space"), space)
        left.set(qn("w:color"), hex_color)
        pBdr.append(left)
        pPr.append(pBdr)

    doc = Document()

    # Seitenränder
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # Standardschrift
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # Kopfzeile
    hdr = doc.sections[0].header
    hdr_p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hdr_p.clear()
    r1 = hdr_p.add_run("AUTODOC")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(9)
    r2 = hdr_p.add_run("  Info Agent · Industry Report")
    r2.font.color.rgb = MUTED; r2.font.size = Pt(9)

    # Fußzeile
    ftr = doc.sections[0].footer
    ftr_p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr_r = ftr_p.add_run(f"{today_str}  ·  AUTODOC Info Agent")
    ftr_r.font.color.rgb = MUTED; ftr_r.font.size = Pt(8)

    # Kategorie-Badge
    bp = doc.add_paragraph()
    br = bp.add_run("AUTODOC · INDUSTRY INTELLIGENCE")
    br.bold = True; br.font.size = Pt(8); br.font.color.rgb = ORANGE
    bp.paragraph_format.space_after = Pt(6)

    # Titel
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(6)
    tr = tp.add_run(data.get("headline", ""))
    tr.bold = True; tr.font.size = Pt(24); tr.font.color.rgb = DARK; tr.font.name = "Arial"

    # Untertitel
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sr = sp.add_run(data.get("subheadline", ""))
    sr.font.size = Pt(13); sr.font.color.rgb = MUTED; sr.font.italic = True

    # Meta
    mp = doc.add_paragraph()
    mp.paragraph_format.space_after = Pt(16)
    mr = mp.add_run(f"AUTODOC Info Agent  ·  {today_str}")
    mr.font.size = Pt(9); mr.font.color.rgb = MUTED

    # Hero-Bild
    hero_img = download_image(data.get("hero_image_url", ""))
    if hero_img:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.add_run().add_picture(hero_img, width=Inches(6.0))
        hp.paragraph_format.space_after = Pt(16)

    # Key-Facts-Tabelle
    key_facts = data.get("key_facts", [])
    if key_facts:
        n = min(len(key_facts), 4)
        kf_tbl = doc.add_table(rows=2, cols=n)
        kf_tbl.style = "Table Grid"
        for i, kf in enumerate(key_facts[:n]):
            val_cell = kf_tbl.rows[0].cells[i]
            set_cell_bg(val_cell, "131C20")
            val_cell.paragraphs[0].clear()
            vr = val_cell.paragraphs[0].add_run(kf.get("value", ""))
            vr.bold = True; vr.font.color.rgb = ORANGE; vr.font.size = Pt(10)
            val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            lbl_cell = kf_tbl.rows[1].cells[i]
            set_cell_bg(lbl_cell, "131C20")
            lbl_cell.paragraphs[0].clear()
            lr = lbl_cell.paragraphs[0].add_run(kf.get("label", "").upper())
            lr.font.color.rgb = PEACH; lr.font.size = Pt(7)
            lbl_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Intro
    for p_text in data.get("intro", "").split("\n"):
        if p_text.strip():
            pp = doc.add_paragraph()
            pr = pp.add_run(p_text.strip())
            pr.font.size = Pt(11)
            pp.paragraph_format.space_after = Pt(8)

    # Pull-Quote
    pq = data.get("pull_quote", "")
    if pq:
        pqp = doc.add_paragraph()
        pqp.paragraph_format.left_indent = Inches(0.4)
        pqp.paragraph_format.space_before = Pt(14)
        pqp.paragraph_format.space_after = Pt(14)
        pqr = pqp.add_run(f'"{pq}"')
        pqr.font.italic = True; pqr.font.size = Pt(13); pqr.font.color.rgb = DARK
        add_para_border_left(pqp)

    # Sektionen
    for section in data.get("sections", []):
        # Überschrift
        h2p = doc.add_paragraph()
        h2p.paragraph_format.space_before = Pt(18)
        h2p.paragraph_format.space_after = Pt(6)
        h2r = h2p.add_run(section.get("heading", ""))
        h2r.bold = True; h2r.font.size = Pt(14); h2r.font.color.rgb = DARK

        # Sektionsbild
        sec_img = download_image(section.get("image_url", ""))
        if sec_img:
            sip = doc.add_paragraph()
            sip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sip.add_run().add_picture(sec_img, width=Inches(6.0))
            sip.paragraph_format.space_after = Pt(10)

        # Fließtext
        for p_text in section.get("content", "").split("\n"):
            if p_text.strip():
                pp = doc.add_paragraph()
                pr = pp.add_run(p_text.strip())
                pr.font.size = Pt(11)
                pp.paragraph_format.space_after = Pt(6)

        # Datentabelle
        if section.get("has_table") and section.get("table"):
            tbl_data = section["table"]
            headers_row = tbl_data.get("headers", [])
            data_rows   = tbl_data.get("rows", [])
            if headers_row and data_rows:
                n_cols = len(headers_row)
                dt = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
                dt.style = "Table Grid"
                for j, h in enumerate(headers_row):
                    c = dt.rows[0].cells[j]
                    set_cell_bg(c, "131C20")
                    c.paragraphs[0].clear()
                    cr = c.paragraphs[0].add_run(h)
                    cr.bold = True; cr.font.color.rgb = WHITE; cr.font.size = Pt(9)
                for ri, row_vals in enumerate(data_rows):
                    bg = "FFFFFF" if ri % 2 == 0 else "F9F9F9"
                    for j, cell_text in enumerate(row_vals):
                        c = dt.rows[ri + 1].cells[j]
                        set_cell_bg(c, bg)
                        c.paragraphs[0].clear()
                        cr = c.paragraphs[0].add_run(str(cell_text))
                        cr.font.size = Pt(9)
                caption = tbl_data.get("caption", "")
                if caption:
                    cp = doc.add_paragraph(caption)
                    cp.runs[0].font.size = Pt(9)
                    cp.runs[0].font.italic = True
                    cp.runs[0].font.color.rgb = MUTED
                doc.add_paragraph()

    # Fazit / Ausblick
    conclusion = data.get("conclusion", "")
    if conclusion:
        conc_label = {"de": "Ausblick", "en": "Outlook", "uk": "Висновок", "ru": "Заключение"}.get(lang, "Ausblick")
        cl = doc.add_paragraph()
        cl.paragraph_format.space_before = Pt(18)
        clr = cl.add_run(conc_label)
        clr.bold = True; clr.font.size = Pt(13); clr.font.color.rgb = DARK
        for p_text in conclusion.split("\n"):
            if p_text.strip():
                pp = doc.add_paragraph()
                pr = pp.add_run(p_text.strip())
                pr.font.size = Pt(11)
                pp.paragraph_format.space_after = Pt(6)

    # Tags
    tags = data.get("tags", [])
    if tags:
        tagp = doc.add_paragraph()
        tagp.paragraph_format.space_before = Pt(16)
        tagr = tagp.add_run("Tags: ")
        tagr.bold = True; tagr.font.size = Pt(9); tagr.font.color.rgb = MUTED
        tagv = tagp.add_run(" · ".join(tags))
        tagv.font.size = Pt(9); tagv.font.color.rgb = MUTED

    # Quellen
    sources = data.get("sources", [])
    if sources:
        src_lbl = {"de": "QUELLEN", "en": "SOURCES", "uk": "ДЖЕРЕЛА", "ru": "ИСТОЧНИКИ"}.get(lang, "QUELLEN")
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(16)
        shr = sh.add_run(src_lbl)
        shr.bold = True; shr.font.size = Pt(9); shr.font.color.rgb = MUTED
        for src in sources:
            srcp = doc.add_paragraph()
            srcp.paragraph_format.space_after = Pt(2)
            arrow_r = srcp.add_run("→ ")
            arrow_r.font.color.rgb = ORANGE; arrow_r.font.size = Pt(9)
            src_name = src.get("name", src.get("url", ""))
            src_url  = src.get("url", "#")
            try:
                r_id = srcp.part.relate_to(
                    src_url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True
                )
                hl = OxmlElement("w:hyperlink")
                hl.set(qn("r:id"), r_id)
                wr = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                col_el = OxmlElement("w:color"); col_el.set(qn("w:val"), "0563C1"); rPr.append(col_el)
                u_el   = OxmlElement("w:u");     u_el.set(qn("w:val"), "single");  rPr.append(u_el)
                sz_el  = OxmlElement("w:sz");    sz_el.set(qn("w:val"), "18");      rPr.append(sz_el)
                wr.append(rPr)
                t_el = OxmlElement("w:t"); t_el.text = src_name; wr.append(t_el)
                hl.append(wr)
                srcp._p.append(hl)
            except Exception:
                srcp.add_run(src_name).font.size = Pt(9)

    # Als Datei-Download zurückgeben
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from flask import send_file as flask_send_file
    headline_raw = data.get("headline", f"artikel_{article_id}")
    safe_name = re.sub(r"[^\w\s-]", "", headline_raw)[:50].strip().replace(" ", "_")
    filename = f"{safe_name}_{row['created_date']}.docx"

    return flask_send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def render_article_html(data: dict, lang: str = "en") -> str:
    """Rendert einen vollständigen, publikationsreifen HTML-Artikel."""
    sections_html = ""
    for i, section in enumerate(data.get("sections", [])):
        img_url = section.get("image_url", "")
        table_html = ""
        if section.get("has_table") and section.get("table"):
            tbl = section["table"]
            headers = "".join(f"<th>{h}</th>" for h in tbl.get("headers", []))
            rows = "".join(
                "<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>"
                for row in tbl.get("rows", [])
            )
            caption = tbl.get("caption", "")
            table_html = f"""
            <div class="table-wrap">
              <table>
                {"<caption>" + caption + "</caption>" if caption else ""}
                <thead><tr>{headers}</tr></thead>
                <tbody>{rows}</tbody>
              </table>
            </div>"""

        img_block = f'<img src="{img_url}" alt="{section.get("heading","")}" class="section-img">' if img_url else ""
        content_paragraphs = "".join(
            f"<p>{p.strip()}</p>" for p in section.get("content", "").split("\n") if p.strip()
        )
        sections_html += f"""
        <section class="article-section">
          {img_block}
          <h2>{section.get("heading","")}</h2>
          {content_paragraphs}
          {table_html}
        </section>"""

    key_facts_html = "".join(f"""
      <div class="kf-item">
        <span class="kf-icon">{kf.get("icon","📌")}</span>
        <span class="kf-value">{kf.get("value","")}</span>
        <span class="kf-label">{kf.get("label","")}</span>
      </div>""" for kf in data.get("key_facts", []))

    tags_html = "".join(
        f'<span class="tag">{tag}</span>' for tag in data.get("tags", []))

    sources_html = "".join(
        f'<li><a href="{s.get("url","#")}" target="_blank">{s.get("name","Quelle")}</a></li>'
        for s in data.get("sources", [])
    )

    today_str = date.today().strftime("%d. %B %Y" if lang == "de" else "%B %d, %Y")
    intro_paragraphs = "".join(
        f"<p>{p.strip()}</p>" for p in data.get("intro", "").split("\n") if p.strip()
    )
    conclusion_paragraphs = "".join(
        f"<p>{p.strip()}</p>" for p in data.get("conclusion", "").split("\n") if p.strip()
    )

    return f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>{data.get("headline","Article")}</title>
  <meta name="description" content="{data.get("meta_description","")}">
  <style>
    :root {{
      --orange: #F85A00; --peach: #FFBA96; --dark: #131C20;
      --text: #1a1a1a; --muted: #666; --border: #e8e8e8;
      --bg: #fafafa;
    }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: 'Georgia','Times New Roman',serif; background: var(--bg); color: var(--text); line-height: 1.75; }}
    a {{ color: var(--orange); text-decoration: none; }}
    a:hover {{ text-decoration: underline; }}

    /* TOP BAR */
    .topbar {{ background: var(--dark); padding: 10px 0; text-align: center; }}
    .topbar-inner {{ max-width: 900px; margin: 0 auto; padding: 0 24px; display: flex; align-items: center; justify-content: space-between; }}
    .brand {{ color: white; font-family: 'Inter',sans-serif; font-size: 13px; display: flex; align-items: center; gap: 10px; }}
    .brand-mark {{ background: var(--orange); color: white; font-weight: 800; padding: 3px 8px; border-radius: 3px; font-size: 11px; }}
    .topbar-date {{ color: var(--peach); font-size: 11px; font-family: 'Inter',sans-serif; }}

    /* HERO */
    .hero {{ position: relative; height: 420px; overflow: hidden; }}
    .hero img {{ width: 100%; height: 100%; object-fit: cover; }}
    .hero-overlay {{ position: absolute; inset: 0; background: linear-gradient(to bottom, rgba(19,28,32,0) 20%, rgba(19,28,32,.85) 100%); }}
    .hero-content {{ position: absolute; bottom: 0; left: 0; right: 0; padding: 32px; max-width: 900px; margin: 0 auto; }}

    /* HEADER */
    .article-header {{ max-width: 900px; margin: 0 auto; padding: 0 24px; }}
    .article-category {{ display: inline-block; background: var(--orange); color: white; font-family: 'Inter',sans-serif; font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; padding: 4px 10px; border-radius: 3px; margin-bottom: 16px; }}
    h1.headline {{ font-size: 2.4rem; font-weight: 700; line-height: 1.2; color: var(--dark); margin-bottom: 12px; }}
    .subheadline {{ font-size: 1.2rem; color: var(--muted); font-weight: 400; margin-bottom: 20px; line-height: 1.4; }}
    .article-meta {{ display: flex; align-items: center; gap: 16px; font-family: 'Inter',sans-serif; font-size: 12px; color: var(--muted); border-top: 1px solid var(--border); border-bottom: 1px solid var(--border); padding: 12px 0; margin-bottom: 28px; }}
    .meta-sep {{ color: var(--border); }}
    .author-note {{ font-style: italic; color: var(--muted); }}

    /* KEY FACTS */
    .key-facts {{ display: grid; grid-template-columns: repeat(4,1fr); gap: 16px; background: var(--dark); padding: 28px 24px; margin-bottom: 40px; }}
    .kf-item {{ text-align: center; }}
    .kf-icon {{ font-size: 24px; display: block; margin-bottom: 6px; }}
    .kf-value {{ display: block; font-family: 'Inter',sans-serif; font-size: 1.4rem; font-weight: 800; color: var(--orange); margin-bottom: 4px; }}
    .kf-label {{ font-family: 'Inter',sans-serif; font-size: 11px; color: var(--peach); text-transform: uppercase; letter-spacing: .5px; }}

    /* ARTICLE BODY */
    .article-body {{ max-width: 900px; margin: 0 auto; padding: 0 24px 48px; }}
    .intro {{ font-size: 1.15rem; line-height: 1.8; margin-bottom: 32px; color: #333; }}
    .intro p {{ margin-bottom: 1em; }}

    /* PULL QUOTE */
    .pull-quote {{ border-left: 5px solid var(--orange); background: rgba(248,90,0,.05); padding: 20px 24px; margin: 36px 0; font-size: 1.2rem; font-style: italic; color: var(--dark); line-height: 1.5; }}

    /* SECTIONS */
    .article-section {{ margin-bottom: 44px; }}
    .article-section h2 {{ font-size: 1.5rem; color: var(--dark); margin-bottom: 16px; padding-bottom: 8px; border-bottom: 2px solid var(--orange); }}
    .article-section p {{ margin-bottom: 1em; font-size: 1rem; }}
    .section-img {{ width: 100%; height: 260px; object-fit: cover; border-radius: 6px; margin-bottom: 20px; }}

    /* TABLE */
    .table-wrap {{ overflow-x: auto; margin: 24px 0; border-radius: 6px; border: 1px solid var(--border); }}
    table {{ width: 100%; border-collapse: collapse; font-family: 'Inter',sans-serif; font-size: 13px; }}
    caption {{ padding: 10px; font-weight: 600; color: var(--muted); font-style: italic; }}
    thead tr {{ background: var(--dark); color: white; }}
    th {{ padding: 12px 14px; text-align: left; font-weight: 600; font-size: 12px; text-transform: uppercase; letter-spacing: .4px; }}
    td {{ padding: 11px 14px; border-bottom: 1px solid var(--border); }}
    tr:nth-child(even) td {{ background: #f9f9f9; }}
    tr:last-child td {{ border-bottom: none; }}
    tr:hover td {{ background: rgba(248,90,0,.04); }}

    /* CONCLUSION */
    .conclusion {{ background: linear-gradient(135deg, rgba(248,90,0,.06), rgba(255,186,150,.08)); border-radius: 8px; padding: 28px; margin: 40px 0; }}
    .conclusion h2 {{ font-size: 1.2rem; margin-bottom: 14px; color: var(--dark); }}
    .conclusion p {{ margin-bottom: .8em; }}

    /* TAGS */
    .tags {{ margin: 32px 0 16px; display: flex; flex-wrap: wrap; gap: 8px; }}
    .tag {{ background: var(--bg); border: 1px solid var(--border); color: var(--muted); font-family: 'Inter',sans-serif; font-size: 11px; padding: 4px 10px; border-radius: 20px; }}

    /* SOURCES */
    .sources {{ border-top: 1px solid var(--border); padding-top: 24px; margin-top: 32px; }}
    .sources h3 {{ font-family: 'Inter',sans-serif; font-size: 12px; text-transform: uppercase; letter-spacing: .8px; color: var(--muted); margin-bottom: 12px; }}
    .sources ul {{ list-style: none; }}
    .sources li {{ font-size: 12px; margin-bottom: 6px; }}
    .sources li::before {{ content: "→ "; color: var(--orange); }}

    /* FOOTER */
    footer {{ background: var(--dark); color: var(--peach); text-align: center; padding: 20px; font-family: 'Inter',sans-serif; font-size: 11px; margin-top: 48px; }}

    @media (max-width: 640px) {{
      .key-facts {{ grid-template-columns: repeat(2,1fr); }}
      h1.headline {{ font-size: 1.6rem; }}
      .hero {{ height: 260px; }}
    }}
    @media print {{
      .topbar {{ display: none; }}
      body {{ background: white; }}
    }}
  </style>
</head>
<body>

<div class="topbar">
  <div class="topbar-inner">
    <div class="brand">
      <span class="brand-mark">AUTODOC</span>
      <span style="color:#aaa">Info Agent · Industry Report</span>
    </div>
    <span class="topbar-date">{today_str}</span>
  </div>
</div>

<div class="hero">
  <img src="{data.get('hero_image_url','')}" alt="{data.get('headline','')}">
  <div class="hero-overlay"></div>
</div>

<div class="article-header" style="padding-top:32px">
  <span class="article-category">AUTODOC · Industry Intelligence</span>
  <h1 class="headline">{data.get("headline","")}</h1>
  <p class="subheadline">{data.get("subheadline","")}</p>
  <div class="article-meta">
    <span>AUTODOC Info Agent</span>
    <span class="meta-sep">·</span>
    <span>{today_str}</span>
    <span class="meta-sep">·</span>
    <span class="author-note">{data.get("author_note","")}</span>
  </div>
</div>

<div class="key-facts">
  {key_facts_html}
</div>

<div class="article-body">

  <div class="intro">
    {intro_paragraphs}
  </div>

  <blockquote class="pull-quote">
    "{data.get("pull_quote","")}"
  </blockquote>

  {sections_html}

  <div class="conclusion">
    <h2>{"Ausblick" if lang == "de" else "Outlook" if lang == "en" else "Висновок" if lang == "uk" else "Заключение"}</h2>
    {conclusion_paragraphs}
  </div>

  <div class="tags">
    {tags_html}
  </div>

  <div class="sources">
    <h3>{"Quellen" if lang == "de" else "Sources" if lang == "en" else "Джерела" if lang == "uk" else "Источники"}</h3>
    <ul>{sources_html}</ul>
  </div>
</div>

<footer>
  © {date.today().year} AUTODOC Info Agent · Generated {today_str} · For internal use
</footer>

</body>
</html>"""

# ── Eigene Quellen (Custom Sources) ──────────────────────────────────────────
@app.route("/api/custom_sources", methods=["GET"])
def get_custom_sources():
    conn = get_db()
    rows = conn.execute("SELECT * FROM custom_sources ORDER BY id DESC").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/custom_sources", methods=["POST"])
def create_custom_source():
    data = request.json
    today = date.today().isoformat()
    if not data.get("name") or not data.get("url"):
        return jsonify({"error": "Name und URL sind Pflichtfelder."}), 400
    conn = get_db()
    import base64
    pw = data.get("password", "")
    pw_b64 = base64.b64encode(pw.encode()).decode() if pw else ""
    conn.execute("""
        INSERT INTO custom_sources (name, url, category, region, relevance, language, active, added_date,
            subscriber_rss, api_key, username, password_b64, cookie)
        VALUES (?, ?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
    """, (data["name"], data["url"],
          data.get("category", "Branchennews"), data.get("region", "GLOBAL"),
          int(data.get("relevance", 5)), data.get("language", "en"), today,
          data.get("subscriber_rss", ""), data.get("api_key", ""),
          data.get("username", ""), pw_b64, data.get("cookie", "")))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/custom_sources/<int:src_id>", methods=["PUT"])
def update_custom_source(src_id):
    data = request.json
    conn = get_db()
    import base64
    pw = data.get("password", "")
    if pw:
        pw_b64 = base64.b64encode(pw.encode()).decode()
    else:
        # Bestehendes Passwort aus DB beibehalten wenn kein neues übermittelt
        existing = conn.execute("SELECT password_b64 FROM custom_sources WHERE id=?", (src_id,)).fetchone()
        pw_b64 = existing["password_b64"] if existing else ""
    conn.execute("""
        UPDATE custom_sources SET name=?, url=?, category=?, region=?,
        relevance=?, language=?, active=?,
        subscriber_rss=?, api_key=?, username=?, password_b64=?, cookie=?
        WHERE id=?
    """, (data.get("name"), data.get("url"),
          data.get("category", "Branchennews"), data.get("region", "GLOBAL"),
          int(data.get("relevance", 5)), data.get("language", "en"),
          int(data.get("active", 1)),
          data.get("subscriber_rss", ""), data.get("api_key", ""),
          data.get("username", ""), pw_b64, data.get("cookie", ""),
          src_id))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/custom_sources/<int:src_id>", methods=["DELETE"])
def delete_custom_source(src_id):
    conn = get_db()
    conn.execute("DELETE FROM custom_sources WHERE id = ?", (src_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# ── Status ────────────────────────────────────────────────────────────────────
@app.route("/api/status")
def api_status():
    today = date.today().isoformat()
    conn = get_db()
    count_today = conn.execute(
        "SELECT COUNT(*) FROM articles WHERE fetched_date = ?", (today,)).fetchone()[0]
    # Fallback auf neuestes Datum wenn heute 0
    if count_today == 0:
        latest = conn.execute("SELECT MAX(fetched_date) FROM articles").fetchone()[0]
        if latest and latest != today:
            count_today = conn.execute(
                "SELECT COUNT(*) FROM articles WHERE fetched_date = ?", (latest,)).fetchone()[0]
    count_total = conn.execute("SELECT COUNT(*) FROM articles").fetchone()[0]
    count_favs = conn.execute("SELECT COUNT(*) FROM articles WHERE is_favorite = 1").fetchone()[0]
    count_paid = conn.execute(
        "SELECT COUNT(*) FROM source_credentials WHERE active = 1").fetchone()[0]
    topics = conn.execute("SELECT COUNT(*) FROM user_topics WHERE active = 1").fetchone()[0]
    conn.close()
    config = load_config()
    return jsonify({
        "articles_today": count_today,
        "articles_total": count_total,
        "favorites": count_favs,
        "paid_sources_active": count_paid,
        "active_topics": topics,
        "has_api_key": bool(config.get("claude_api_key", "") not in ["", "DEIN_API_KEY_HIER"]),
        "language": get_setting("language", "de"),
        "relevance_min": int(get_setting("relevance_min", "1")),
        "date": today,
        "sources_total": 100,
        "feed_running": feed_status.get("running", False),
        "last_fetch": feed_status.get("last_fetch", "—"),
        "feed_new_count": feed_status.get("new_count", 0)
    })

# ── Kategorien ────────────────────────────────────────────────────────────────
@app.route("/api/categories")
def api_categories():
    today = date.today().isoformat()
    conn = get_db()
    count_today = conn.execute(
        "SELECT COUNT(*) FROM articles WHERE fetched_date = ?", (today,)).fetchone()[0]
    if count_today == 0:
        latest = conn.execute("SELECT MAX(fetched_date) FROM articles").fetchone()[0]
        display_date = latest if latest else today
    else:
        display_date = today
    rows = conn.execute("""
        SELECT category, COUNT(*) as count FROM articles
        WHERE fetched_date = ?
        GROUP BY category ORDER BY count DESC
    """, (display_date,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ── Kategorie-Einstellungen ───────────────────────────────────────────────────
@app.route("/api/category_settings", methods=["GET"])
def get_category_settings():
    """Gibt alle Kategorien mit ihrer User-Priorität zurück."""
    conn = get_db()
    # Alle tatsächlich vorhandenen Kategorien aus Artikeln + vorgespeicherte
    cats_in_articles = [r["category"] for r in
        conn.execute("SELECT DISTINCT category FROM articles WHERE category IS NOT NULL AND category != '' ORDER BY category").fetchall()]
    existing = {r["category"]: dict(r) for r in
        conn.execute("SELECT * FROM category_settings").fetchall()}
    conn.close()

    all_cats = sorted(set(cats_in_articles) | set(existing.keys()))
    result = []
    for cat in all_cats:
        if cat:
            result.append({
                "category": cat,
                "priority": existing.get(cat, {}).get("priority", 5),
                "active": existing.get(cat, {}).get("active", 1)
            })
    return jsonify(result)

@app.route("/api/category_settings", methods=["POST"])
def save_category_settings():
    """Speichert Kategorie-Prioritäten."""
    settings = request.json.get("settings", [])
    conn = get_db()
    for s in settings:
        if s.get("category"):
            conn.execute("""INSERT OR REPLACE INTO category_settings (category, priority, active)
                           VALUES (?, ?, ?)""",
                        (s["category"], int(s.get("priority", 5)), int(s.get("active", 1))))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# ── Start ─────────────────────────────────────────────────────────────────────
APP_HOST = "autodoc-info-agent"
APP_PORT = 5050

def get_app_url():
    """Gibt die beste verfügbare URL zurück."""
    import socket
    try:
        socket.gethostbyname(APP_HOST)
        return f"http://{APP_HOST}:{APP_PORT}"
    except Exception:
        return f"http://localhost:{APP_PORT}"

if __name__ == "__main__":
    print("=" * 55)
    print("  AUTODOC Info Agent v3 wird gestartet...")
    print("=" * 55)
    init_db()

    refresh_hour = int(get_setting("refresh_hour", "7"))
    scheduler = BackgroundScheduler()
    scheduler.add_job(fetch_feeds, "cron", hour=refresh_hour, minute=0)
    scheduler.start()

    threading.Thread(target=fetch_feeds, daemon=True).start()

    app_url = get_app_url()

    def open_browser():
        time.sleep(2)
        webbrowser.open(app_url)
    threading.Thread(target=open_browser, daemon=True).start()

    print(f"  Browser öffnet sich in 2 Sekunden...")
    print(f"  ➜  {app_url}")
    print("  Strg+C zum Beenden")
    print("=" * 55)
    app.run(host="0.0.0.0", port=APP_PORT, debug=False)
